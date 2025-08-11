#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
将Markdown文档转换为Word文档的脚本
包含目录和跳转功能
"""

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.shared import OxmlElement, qn
import re

def add_hyperlink(paragraph, text, url):
    """添加超链接"""
    part = paragraph.part
    r_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
    
    # Create the w:hyperlink tag
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)
    
    # Create a w:r element
    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    
    # Join all the xml elements together
    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)
    
    # Create a new Run object and add the hyperlink into it
    r = paragraph.add_run()
    r._r.append(hyperlink)
    
    # A workaround for the lack of a hyperlink style (doesn't go purple after using)
    r.font.color.theme_color = MSO_THEME_COLOR_INDEX.HYPERLINK
    r.font.underline = True
    
    return hyperlink

def create_word_document():
    """创建Word文档"""
    doc = Document()
    
    # 设置文档标题
    title = doc.add_heading('金融数据处理工具函数文档', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 添加目录标题
    toc_heading = doc.add_heading('目录', level=1)
    
    # 读取Markdown文件
    with open('function_documentation.md', 'r', encoding='utf-8') as f:
        content = f.read()
    
    # 解析内容
    sections = content.split('## 函数详细说明')[1]
    
    # 添加utils.py函数
    doc.add_heading('utils.py 函数', level=2)
    
    # 添加函数列表
    utils_functions = [
        'source_getting', 'source_getting2', 'get_db_connection', 'close_all_connections',
        'data_reader', 'data_getting_glb', 'data_getting', 'contains_chinese', 'index_mapping',
        'readcsv', 'chunks', 'file_withdraw', 'file_withdraw2', 'folder_creator',
        'folder_creator2', 'folder_creator3', 'move_specific_files', 'move_specific_files2',
        'get_string_before_last_dot', 'optiondata_greeksprocessing'
    ]
    
    for func in utils_functions:
        p = doc.add_paragraph()
        p.add_run(f'• {func}').bold = True
    
    # 添加time_utils.py函数
    doc.add_heading('time_utils.py 函数', level=2)
    
    time_functions = [
        'Chinese_valuation_date', 'next_workday_auto', 'last_workday_auto',
        'last_workday_calculate', 'next_workday_calculate', 'last_workday_calculate2',
        'is_workday', 'working_days', 'is_workday_auto', 'intdate_transfer',
        'strdate_transfer', 'working_days_list', 'working_day_count',
        'month_lastday_df', 'last_weeks_lastday_df', 'last_weeks_lastday',
        'weeks_firstday', 'next_weeks_lastday'
    ]
    
    for func in time_functions:
        p = doc.add_paragraph()
        p.add_run(f'• {func}').bold = True
    
    # 添加global_tools.py函数
    doc.add_heading('global_tools.py 函数', level=2)
    
    global_functions = [
        'sql_to_timeseries', 'rank_score_processing', 'code_transfer',
        'factor_name_old', 'factor_name_new', 'factor_name', 'weight_sum_check',
        'weight_sum_warning', 'stock_volatility_calculate', 'factor_universe_withdraw',
        'index_weight_withdraw', 'indexData_withdraw', 'indexFactor_withdraw',
        'stockData_withdraw', 'etfData_withdraw', 'cbData_withdraw',
        'optionData_withdraw', 'futureData_withdraw', 'weight_df_standardization',
        'weight_df_datecheck', 'portfolio_analyse', 'backtesting_report', 'table_manager'
    ]
    
    for func in global_functions:
        p = doc.add_paragraph()
        p.add_run(f'• {func}').bold = True
    
    # 添加详细说明
    doc.add_heading('函数详细说明', level=1)
    
    # 解析函数详细说明
    function_details = re.findall(r'#### (\w+)\n\*\*函数名\*\*: (\w+)\s*\n\*\*输入\*\*: (.*?)\n\*\*输出\*\*: (.*?)\n\*\*说明\*\*: (.*?)(?=\n####|\n---|\Z)', 
                                 sections, re.DOTALL)
    
    for func_name, func_id, inputs, outputs, description in function_details:
        # 添加函数标题
        heading = doc.add_heading(func_name, level=2)
        
        # 添加函数信息表格
        table = doc.add_table(rows=4, cols=2)
        table.style = 'Table Grid'
        
        # 设置表格内容
        table.cell(0, 0).text = '函数名'
        table.cell(0, 1).text = func_id
        
        table.cell(1, 0).text = '输入'
        table.cell(1, 1).text = inputs.strip()
        
        table.cell(2, 0).text = '输出'
        table.cell(2, 1).text = outputs.strip()
        
        table.cell(3, 0).text = '说明'
        table.cell(3, 1).text = description.strip()
        
        # 设置表格样式
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        
        doc.add_paragraph()  # 添加空行
    
    # 添加使用示例
    doc.add_heading('使用示例', level=1)
    
    # 添加数据获取示例
    doc.add_heading('数据获取示例', level=2)
    example1 = doc.add_paragraph()
    example1.add_run('获取股票数据：\n')
    example1.add_run('df_stock = stockData_withdraw("2024-01-01", "2024-01-31", ["close", "pre_close"])\n\n')
    example1.add_run('获取指数权重：\n')
    example1.add_run('df_weight = index_weight_withdraw("沪深300", "2024-01-31")\n\n')
    example1.add_run('获取交易日期：\n')
    example1.add_run('df_date = Chinese_valuation_date()')
    
    # 添加时间处理示例
    doc.add_heading('时间处理示例', level=2)
    example2 = doc.add_paragraph()
    example2.add_run('获取下一个工作日：\n')
    example2.add_run('next_day = next_workday_auto()\n\n')
    example2.add_run('判断是否为工作日：\n')
    example2.add_run('is_work = is_workday("2024-01-31")\n\n')
    example2.add_run('获取工作日列表：\n')
    example2.add_run('work_days = working_days_list("2024-01-01", "2024-01-31")')
    
    # 添加投资组合分析示例
    doc.add_heading('投资组合分析示例', level=2)
    example3 = doc.add_paragraph()
    example3.add_run('投资组合分析：\n')
    example3.add_run('df_info, df_detail = portfolio_analyse(\n')
    example3.add_run('    df_holding=df_holding,\n')
    example3.add_run('    account_money=10000000,\n')
    example3.add_run('    cost_stock=0.00085,\n')
    example3.add_run('    cost_etf=0.0003\n')
    example3.add_run(')')
    
    # 添加注意事项
    doc.add_heading('注意事项', level=1)
    notes = doc.add_paragraph()
    notes.add_run('1. 所有数据获取函数都支持本地文件和数据库两种模式\n')
    notes.add_run('2. 时间相关函数基于中国交易日历\n')
    notes.add_run('3. 投资组合分析函数需要完整的市场数据支持\n')
    notes.add_run('4. 数据库连接使用连接池管理，提高性能\n')
    notes.add_run('5. 文件操作函数支持多种编码格式\n')
    notes.add_run('6. 权重标准化函数支持股票、期货、期权等多种资产类型')
    
    # 添加版本信息
    doc.add_heading('版本信息', level=1)
    version = doc.add_paragraph()
    version.add_run('• 文档版本: 1.0\n')
    version.add_run('• 创建日期: 2024年\n')
    version.add_run('• 适用Python版本: 3.7+\n')
    version.add_run('• 主要依赖: pandas, numpy, pymysql, scipy')
    
    # 保存文档
    doc.save('金融数据处理工具函数文档.docx')
    print("Word文档已生成：金融数据处理工具函数文档.docx")

if __name__ == "__main__":
    create_word_document()

